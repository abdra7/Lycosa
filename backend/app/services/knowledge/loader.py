"""Document loading and chunking: bytes in, text chunks out."""

import csv
import io
import json
from io import BytesIO
from typing import Any


class ExtractionError(ValueError):
    """The document's text could not be extracted; message is operator-facing."""


def _extract_pdf(filename: str, data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            # some PDFs carry an owner password only; an empty user password opens them
            try:
                decrypted = bool(reader.decrypt(""))
            except Exception:
                decrypted = False
            if not decrypted:
                raise ExtractionError(
                    f"PDF {filename!r} is password-protected — upload a decrypted copy"
                )
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(
            f"could not parse PDF {filename!r} (corrupt or unsupported file): {exc}"
        ) from exc


def _extract_docx(filename: str, data: bytes) -> str:
    """Paragraphs become blank-line-separated blocks (so the paragraph chunker
    packs them); each table row becomes a 'cell | cell | ...' block."""
    import docx

    try:
        document = docx.Document(BytesIO(data))
    except Exception as exc:
        raise ExtractionError(
            f"could not parse DOCX {filename!r} (corrupt or unsupported file): {exc}"
        ) from exc
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    if not blocks:
        raise ExtractionError(f"DOCX {filename!r} contained no extractable text")
    return "\n\n".join(blocks)


def _extract_csv(filename: str, data: bytes) -> str:
    """Each row becomes a 'header: value | ...' paragraph so retrieval can
    isolate a single record. The first row is taken as the header (ADR-024)."""
    text = data.decode("utf-8-sig", errors="replace")  # -sig strips a BOM if present
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error as exc:
        raise ExtractionError(f"could not parse CSV {filename!r}: {exc}") from exc
    rows = [r for r in rows if any(cell.strip() for cell in r)]  # drop blank lines
    if len(rows) < 2:
        raise ExtractionError(f"CSV {filename!r} has no data rows below the header")
    header = [h.strip() for h in rows[0]]
    records = []
    for row in rows[1:]:
        fields = []
        for i, value in enumerate(row):
            value = value.strip()
            if not value:
                continue
            label = header[i] if i < len(header) and header[i] else f"col{i + 1}"
            fields.append(f"{label}: {value}")
        if fields:
            records.append(" | ".join(fields))
    if not records:
        raise ExtractionError(f"CSV {filename!r} has no non-empty data rows")
    return "\n\n".join(records)


def _json_scalar(value: Any) -> str:
    if value is None:
        return "null"
    if value is True:
        return "true"
    if value is False:
        return "false"
    return str(value)


def _flatten_json(node: Any, prefix: str = "") -> list[str]:
    """Flatten to 'a.b[0]: value' leaf lines, keeping structure searchable."""
    if isinstance(node, dict):
        lines: list[str] = []
        for key, value in node.items():
            lines.extend(_flatten_json(value, f"{prefix}.{key}" if prefix else str(key)))
        return lines
    if isinstance(node, list):
        lines = []
        for index, value in enumerate(node):
            lines.extend(_flatten_json(value, f"{prefix}[{index}]"))
        return lines
    leaf = _json_scalar(node)
    return [f"{prefix}: {leaf}" if prefix else leaf]


def _extract_json(filename: str, data: bytes) -> str:
    """A top-level array yields one record (paragraph) per element; any other
    value is flattened to path/value lines (ADR-024)."""
    try:
        parsed = json.loads(data.decode("utf-8-sig", errors="replace"))
    except (json.JSONDecodeError, ValueError) as exc:
        raise ExtractionError(f"could not parse JSON {filename!r}: {exc}") from exc
    if isinstance(parsed, list):
        records = ["\n".join(_flatten_json(item)) for item in parsed]
        records = [r for r in records if r.strip()]
    else:
        records = ["\n".join(_flatten_json(parsed))]
    text = "\n\n".join(r for r in records if r.strip())
    if not text.strip():
        raise ExtractionError(f"JSON {filename!r} contained no extractable values")
    return text


_ZIP_MAGIC = b"PK\x03\x04"  # docx/xlsx/pptx/odt/zip
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"  # legacy .doc/.xls/.ppt
# A real text file may carry a few mojibake bytes (still replaced, not fatal);
# wholesale binary decodes to a large share of replacement chars — refuse it.
_MAX_REPLACEMENT_CHARS = 4
_MAX_REPLACEMENT_RATIO = 0.05


def _decode_text(filename: str, data: bytes) -> str:
    """UTF-8 fallback for md/txt/code, refusing binary content: silently
    embedding undecodable bytes stores unsearchable junk as 'embedded' (#28)."""
    if data.startswith(_ZIP_MAGIC) or data.startswith(_OLE2_MAGIC):
        raise ExtractionError(
            f"{filename!r} is a binary document format with no loader — "
            "supported formats: .pdf, .docx, .csv, .json, and plain text"
        )
    text = data.decode("utf-8", errors="replace")
    replacements = text.count("�")
    if "\x00" in text or (
        replacements > _MAX_REPLACEMENT_CHARS and replacements / len(text) > _MAX_REPLACEMENT_RATIO
    ):
        raise ExtractionError(
            f"{filename!r} is not UTF-8 text — "
            "supported formats: .pdf, .docx, .csv, .json, and plain text"
        )
    return text


def extract_text(filename: str, data: bytes) -> str:
    """PDF via pypdf; DOCX via python-docx; CSV/JSON parsed structure-aware
    (ADR-024); everything else must decode as UTF-8 text (md/txt/code)."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(filename, data)
    if name.endswith(".docx"):
        return _extract_docx(filename, data)
    if name.endswith(".csv"):
        return _extract_csv(filename, data)
    if name.endswith(".json"):
        return _extract_json(filename, data)
    return _decode_text(filename, data)


def chunk_text(text: str, size: int = 800, overlap: int = 100) -> list[str]:
    """Paragraph-aware chunking: pack paragraphs up to `size` chars; oversized
    paragraphs are hard-split with `overlap` chars of continuity."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""

    def flush() -> None:
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for paragraph in paragraphs:
        if len(paragraph) > size:
            flush()
            start = 0
            while start < len(paragraph):
                chunks.append(paragraph[start : start + size].strip())
                start += size - overlap
        elif len(current) + len(paragraph) + 2 > size:
            flush()
            current = paragraph
        else:
            current = f"{current}\n\n{paragraph}" if current else paragraph
    flush()
    return chunks
