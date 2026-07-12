"""DOCX extraction + binary-content guard (issue #28).

A .docx used to fall through to the UTF-8 fallback: the OOXML ZIP bytes were
decoded as replacement-character garbage and reported as successfully embedded.
Now .docx is parsed with python-docx, and the plain-text fallback refuses
binary content (ZIP/OLE2 magic, NUL bytes, high replacement-char ratio) with an
ExtractionError so no binary format is ever silently ingested as junk.
"""

from io import BytesIO

import pytest
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import DocumentStatus, KnowledgeCollection
from app.services.knowledge.embedder import get_embedder
from app.services.knowledge.ingestion import ingest_document
from app.services.knowledge.loader import ExtractionError, extract_text


def _build_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        docx_table = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                docx_table.cell(r, c).text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# .docx extraction
# ---------------------------------------------------------------------------


def test_docx_paragraphs_become_blank_line_separated_text() -> None:
    data = _build_docx(
        [
            "New drivers complete a two-day safety orientation before their first route.",
            "Wolf spiders hunt at night without webs.",
        ]
    )
    text = extract_text("onboarding.docx", data)
    paragraphs = text.split("\n\n")
    assert paragraphs[0].startswith("New drivers complete a two-day safety orientation")
    assert "Wolf spiders hunt at night" in paragraphs[1]
    # the old bug: ZIP bytes decoded to replacement-char junk
    assert "�" not in text


def test_docx_table_cells_are_extracted() -> None:
    data = _build_docx(["Intro."], table=[["name", "legs"], ["tarantula", "8"]])
    text = extract_text("spiders.docx", data)
    assert "tarantula" in text
    assert "8" in text


def test_docx_extension_is_case_insensitive() -> None:
    data = _build_docx(["Case check."])
    assert "Case check." in extract_text("REPORT.DOCX", data)


def test_docx_with_no_text_raises_extraction_error() -> None:
    with pytest.raises(ExtractionError, match="empty.docx"):
        extract_text("empty.docx", _build_docx([]))


def test_corrupt_docx_raises_extraction_error_naming_the_file() -> None:
    with pytest.raises(ExtractionError, match="broken.docx"):
        extract_text("broken.docx", b"PK\x03\x04 this is not a real zip archive")


# ---------------------------------------------------------------------------
# binary guard on the plain-text fallback
# ---------------------------------------------------------------------------


def test_zip_based_format_without_a_loader_is_rejected_not_garbled() -> None:
    # an .xlsx is a ZIP too; there is no loader for it, so it must fail clearly
    with pytest.raises(ExtractionError, match="book.xlsx"):
        extract_text("book.xlsx", b"PK\x03\x04" + b"\x00" * 64)


def test_legacy_ole2_office_format_is_rejected() -> None:
    with pytest.raises(ExtractionError, match="old.doc"):
        extract_text("old.doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64)


def test_undecodable_binary_txt_is_rejected() -> None:
    with pytest.raises(ExtractionError, match="blob.bin"):
        extract_text("blob.bin", b"\xff\xfe\xfa" * 100)


def test_nul_bytes_are_rejected_as_binary() -> None:
    # UTF-16 bytes decode as "valid" UTF-8 full of NULs — that is not text we
    # can chunk meaningfully, so it must be refused, not embedded
    with pytest.raises(ExtractionError, match="utf16.txt"):
        extract_text("utf16.txt", "hello world".encode("utf-16"))


def test_mostly_valid_utf8_with_a_stray_byte_is_still_accepted() -> None:
    data = ("wolf spiders hunt at night " * 20).encode() + b"\xff"
    text = extract_text("notes.txt", data)
    assert "wolf spiders hunt at night" in text


def test_plain_markdown_still_decoded_as_text() -> None:
    assert extract_text("notes.md", b"# Spiders\n\nhunt at night") == "# Spiders\n\nhunt at night"


# ---------------------------------------------------------------------------
# end-to-end: a docx ingests into retrievable text, not 50 chunks of junk
# ---------------------------------------------------------------------------


async def test_docx_ingests_as_clean_text_chunks(db_session: AsyncSession, qdrant) -> None:
    embedder = get_embedder("hashing")
    collection = KnowledgeCollection(
        name="docx-e2e", embedding_backend=embedder.name, embedding_dim=embedder.dim
    )
    db_session.add(collection)
    await db_session.commit()
    await db_session.refresh(collection)

    data = _build_docx(
        [
            "New drivers complete a two-day safety orientation before their first route.",
            "Routes are assigned every Monday morning.",
        ]
    )
    document = await ingest_document(
        db_session,
        collection,
        "onboarding.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data,
    )

    assert document.status == DocumentStatus.EMBEDDED
    # the regression produced 50 junk chunks from a 2-paragraph file
    assert document.chunk_count == 1


async def test_unsupported_binary_lands_failed_with_clear_error(
    db_session: AsyncSession, qdrant
) -> None:
    embedder = get_embedder("hashing")
    collection = KnowledgeCollection(
        name="binary-guard-e2e", embedding_backend=embedder.name, embedding_dim=embedder.dim
    )
    db_session.add(collection)
    await db_session.commit()
    await db_session.refresh(collection)

    document = await ingest_document(
        db_session, collection, "book.xlsx", None, b"PK\x03\x04" + b"\x00" * 64
    )

    assert document.status == DocumentStatus.FAILED
    assert document.error is not None
    assert "book.xlsx" in document.error
